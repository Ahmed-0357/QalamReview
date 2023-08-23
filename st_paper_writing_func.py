import ast
import os
import time

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from langchain import LLMChain
from langchain.prompts.chat import (ChatPromptTemplate,
                                    HumanMessagePromptTemplate,
                                    SystemMessagePromptTemplate)


def load_and_process_df(file_path, numeric=True):
    """Load and process a relevance and summary files into a DataFrame.

    Args:
        file_path (str): file path
        numeric (bool, optional): convert columns to numeric value. Defaults to True.

    Returns:
        dataframe: processed dataframe
    """

    with open(file_path, 'r', encoding='utf8') as file:
        content = file.read()

    parts = content.split('\n\n')[:-1]
    dicts = [ast.literal_eval(part) for part in parts]

    if numeric:
        df = pd.DataFrame(dicts)
        df = df.apply(pd.to_numeric, errors='coerce')
    else:
        merged_dicts = [{**d['metadata'], **d['summary']} for d in dicts]
        df = pd.DataFrame(merged_dicts)
        df = df.replace("None", "")

    return df


class PaperSummary:
    """summarizes a scholarly paper
    """

    def __init__(self, texts):
        """instantiation of paper summary class

        Args:
            texts (list): list of paper texts
        """
        self.texts = texts
        self.summarize_format = {
            "metadata": {
                "title": "",
                "author(s)": "",
                "journal": "",
                "year": ""},
            "summary": {
                "introduction": "",
                "methodology": "",
                "results/findings": "",
                "limitations/gaps": ""}
        }

    def paper_summary_prompt(self, p_type='once'):
        """outlines creation prompt
        Args:
            p_type(str): type of summary can be either once, many1 or many2. many1&2 for long paper that can not summarized via the model 
        Returns:
            str: prompt string
        """

        if p_type == 'once':
            system_message_prompt = SystemMessagePromptTemplate.from_template("""Possessing a notable reputation as a researcher, and having exceptional skill in dissecting scholarly research papers, along with deep expertise in the field of {expertise_areas},
            your task is to meticulously analyze the provided document. Your goal is to craft a succinct summary of this intricate information, while preserving the integrity, accuracy, and precision of key concepts from the original academic material.
            
            Keep in mind that this summary will be a crucial part of the narrative review paper on the subject of {subject} and structured following the provided outline {outline}. During the construction of the summary, ensure diligent cross-referencing with the outline. If any information aligns with a section from the outline and is discussed or mentioned in the paper, it is crucial to incorporate it into the
            corresponding section of the summary in an appropriate manner.""")

            human_message_prompt = HumanMessagePromptTemplate.from_template("""Using the comprehensive content of the scholarly paper {paper_content}, your task is to distill, analyze, and categorize the information into two key structured sections: Metadata and Summary.
            
            Under "metadata", gather the following details
            
            "title": Identify the title of the paper.
            "author(s)": List the author(s) of the paper.
            "journal": Provide the name of the journal where the paper was published.
            "year": Mention the year of publication.
            
            If any of these details are absent from the paper content, please denote the corresponding field as 'None'.
            
            Under "summary", break down the information into the following categories:

            "introduction": Decode the background, problem statement, primary objectives, and motivations of the study from the paper content. If the content doesn't provide enough information for this, please denote this section as 'None'.

            "methodology": Develop a thorough understanding of the research methodologies used, drawing from the paper content. This might include data collection and analysis strategies, study design details, sample size, experimental setup, any simulations conducted, and analytical tools employed. If such details are not evident from the paper content, label this section as 'None'.

            "results/findings": Extract key findings, pivotal conclusions, and significant data trends or patterns from the paper content. If these details aren't sufficiently outlined, please mark this section as 'None'.

            "limitations/gaps": Identify potential limitations or gaps in the study, as suggested by the paper content. This could involve issues with the study design or research areas left unaddressed. If these aspects aren't clear, please denote this section as 'None'.

            Please assemble and structure the extracted information into a Python dictionary, adhering to the following format: {summarize_format}, also just give the python dictionary without anything else""")

        elif p_type == 'many1':
            system_message_prompt = SystemMessagePromptTemplate.from_template(
                """you have well known academic researcher specialized knowledge in {expertise_areas}""")

            human_message_prompt = HumanMessagePromptTemplate.from_template(
                """Could you please transform the academic article {paper_content} into a series of key points. Strive to retain as much information and details as you can and do not summarize. For the first key point, please incorporate the paper's title, authors, the publishing journal, and the date of publication""")

        else:
            system_message_prompt = SystemMessagePromptTemplate.from_template(
                """you have well known academic researcher specialized knowledge in {expertise_areas}""")

            human_message_prompt = HumanMessagePromptTemplate.from_template(
                """Could you please transform the academic article {paper_content} into a series of key points. Strive to retain as much information and details as you can and do not summarize. """)

        chat_prompt = ChatPromptTemplate.from_messages(
            [system_message_prompt, human_message_prompt])

        return chat_prompt

    def summarize(self, **kwargs):
        """summarize the scholarly paper

        Returns:
            dict: paper summary in the specified format
        """

        if len(self.texts) == 1:
            chat_prompt = self.paper_summary_prompt(p_type='once')
            chain = LLMChain(llm=kwargs.get('llm_model'), prompt=chat_prompt)

            output = chain.run(expertise_areas=kwargs.get('expertise_areas'), subject=kwargs.get('subject'), outline=kwargs.get('outline'),
                               paper_content=self.texts[0], summarize_format=self.summarize_format)
        else:
            section_summary = {}
            for i in range(len(self.texts)):
                if i == 0:
                    chat_prompt = self.paper_summary_prompt(p_type='many1')
                    chain = LLMChain(llm=kwargs.get(
                        'llm_model'), prompt=chat_prompt)
                    r = chain.run(expertise_areas=kwargs.get('expertise_areas'),
                                  paper_content=self.texts[i])
                    section_summary[i] = r
                else:
                    chat_prompt = self.paper_summary_prompt(p_type='many2')
                    chain = LLMChain(llm=kwargs.get(
                        'llm_model'), prompt=chat_prompt)
                    r = chain.run(expertise_areas=kwargs.get('expertise_areas'),
                                  paper_content=self.texts[i])
                    section_summary[i] = r

                    time.sleep(9)

            chat_prompt = self.paper_summary_prompt(p_type='once')
            chain = LLMChain(llm=kwargs.get('llm_model'), prompt=chat_prompt)
            all_sections = ' '.join(list(section_summary.values()))
            output = chain.run(expertise_areas=kwargs.get('expertise_areas'), subject=kwargs.get('subject'), outline=kwargs.get('outline'),
                               paper_content=all_sections, summarize_format=self.summarize_format)

        output_dict = ast.literal_eval(output)
        return output_dict


class RelevanceAnalysis:
    """"Assign a relevancy rating (0-100) to each paper summary, guided by the outline" 
    """

    def __init__(self, paper_summary):
        """instantiate the class

        Args:
            paper_summary (dict): paper summary
        """
        self.paper_summary = str(paper_summary)

    def relevancy_analysis_prompt(self):
        """relevancy prompt
        Returns:
            str: prompt string
        """
        system_message_prompt = SystemMessagePromptTemplate.from_template(
            """Equipped with substantial expertise in {expertise_areas}, you're an accomplished researcher renowned for crafting insightful narrative review papers and conducting thorough assessments of journal paper summaries.
            Your task harnesses your unique aptitude for aligning the essence of a journal summary with the framework of a narrative review paper. Apply your knowledge and experience to successfully complete this assignment.""")
        human_message_prompt = HumanMessagePromptTemplate.from_template(
            """Given a summary of a specific research paper {paper_summary} and the outline of a narrative review paper {outline}, titled {subject}, your task is to evaluate the relevance of the paper's content in relation to each and every section of the review paper's outline. Relevance should be scored from 0 to 100, where 100 indicates strong relevance and 0 signifies no relevance. Once evaluated, organize your results into a Python dictionary. Make sure to strictly follow the provided format {relevance_format} when compiling your results. also just give the python dictionary without anything else""")
        chat_prompt = ChatPromptTemplate.from_messages(
            [system_message_prompt, human_message_prompt])

        return chat_prompt

    def data_parser_prompt(self):
        """parsing prompt
        Returns:
            str: prompt string
        """
        system_message_prompt = SystemMessagePromptTemplate.from_template(
            """you are an expert data converter AI that can convert provided text into JSON format""")
        human_message_prompt = HumanMessagePromptTemplate.from_template(
            """I have the following text {text} that I would like to be converted into JSON following this format {relevance_format}, also just give the JSON without anything else""")
        chat_prompt = ChatPromptTemplate.from_messages(
            [system_message_prompt, human_message_prompt])

        return chat_prompt

    def relevancy_score(self, **kwargs):
        """get relevancy score using llm model. relevancy score ranges between 0 to 100

        Returns:
            dict: relevancy score
        """
        relevance_format = {
            i: "relevance score" for i in kwargs.get('outline')}
        chat_prompt = self.relevancy_analysis_prompt()
        chain = LLMChain(llm=kwargs.get('llm_model')[0], prompt=chat_prompt)
        output = chain.run(expertise_areas=kwargs.get('expertise_areas'), paper_summary=self.paper_summary, outline=kwargs.get(
            'outline'), subject=kwargs.get('subject'), relevance_format=relevance_format)

        # parsing results
        chat_prompt = self.data_parser_prompt()
        chain = LLMChain(llm=kwargs.get('llm_model')[1], prompt=chat_prompt)
        output = chain.run(text=output, relevance_format=relevance_format)

        output_dict = ast.literal_eval(output)
        return output_dict


class ManuscriptWriting:
    """Manuscript writing using the papers summary and relevancy score
    """

    def __init__(self, manuscript_name='narrative_review.docx', manuscript_dir='manuscript'):
        """instantiate the class

        Args:
            manuscript_name (str, optional): manuscript paper name. Defaults to 'narrative_review.docx'.
            manuscript_dir (str, optional): dir to store the manuscript. Defaults to 'manuscript'.
        """
        self.writing_format = {'KK': 'content'}
        # make dir
        self.manuscript_dir = manuscript_dir
        os.makedirs(self.manuscript_dir, exist_ok=True)
        # create docx file
        self.manuscript_name = manuscript_name
        self.file_path = os.path.join(
            self.manuscript_dir, self.manuscript_name)
        self.doc = Document()
        self.doc.save(self.file_path)

    def add_text(self, text, style='Normal', bold=False, italic=False, size=11, centered=False):
        """Add text to the Word document.

        Args:
            text (text): gpt content
            style (str, optional): can be Normal, Heading1, Heading2. Defaults to 'Normal'.
            bold (bool, optional): make font bold. Defaults to False.
            italic (bool, optional): make font italic. Defaults to False.
            size (int, optional): font size. Defaults to 11.
            centered (bool, optional): center the content. Defaults to False.
        """
        paragraph = self.doc.add_paragraph(style=style)
        paragraph.style = self.doc.styles[style]
        run = paragraph.add_run(text)
        run.bold, run.italic, run.font.size = bold, italic, Pt(size)
        if centered:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.doc.save(self.file_path)

    def final_writeup(self, title, data):
        """
        This function processes a nested dictionary of the paper content and adds its content to a Word document.

        Args:
            data (dict): The nested dictionary to process.
            title (str): The title of the document.
        """

        self.add_text(title, 'Normal', bold=True, centered=True, size=16)

        for key1, value1 in data.items():
            self.add_text(f'{key1}', 'Heading 1', bold=True, size=12)
            for key2, value2 in value1.items():
                self.add_text(f'{key2}', 'Heading 2', bold=True, size=12)
                for key3, value3 in value2.items():
                    self.add_text(f'{key3}', 'Heading 3',
                                  bold=True, italic=True, size=11)
                    self.add_text(value3, 'Normal', size=11)

    def section_writing_prompt(self, p_type='short'):
        """section writing prompt
        Args:
            p_type(str): type of writing prompt, default short 
        Returns:
            str: prompt string
        """

        if p_type == 'short':
            system_message_prompt = SystemMessagePromptTemplate.from_template("""As a reputable researcher with comprehensive expertise in {expertise_areas}, you have consistently demonstrated a unique capability for critically evaluating scholarly articles. We anticipate your valuable input in our narrative review paper titled {subject}, utilizing these skills.
                                                                    
            Your skill set extends to the examination and precise extraction of content from scholarly papers, specifically targeting information that aligns with section {section} of our narrative review. You possess a remarkable knack for synthesizing complex data, pinpointing connections, and detecting any discrepancies within the culled data.
            Please ensure that the final output is presented as one cohesive paragraph, with in-text citations for each referenced paper. Note that there's no need to include a separate reference list.
            """)

            human_message_prompt = HumanMessagePromptTemplate.from_template("""Please review the following summaries of scholarly papers: {papers_summary}
                                                                    
            Your task is centered on the extraction, analysis, and synthesis of content that aligns specifically with the section titled "{section}". Instead of utilizing all the content from each paper, your objective is to distill only the information that is strongly relevant to this {section}. Use this distilled information to form a comprehensive narrative, weaving connections and pinpointing discrepancies among the studies.
            
            If the summaries of the papers do not contain strongly relevant content to the {section} or are empty, please return 'None'.
            
            Remember to use APA style for in-text citations accurately for each paper as you merge these insights into the narrative review.
            
            The final output should strictly be a cohesive paragraph with appropriate in-text citations, and there's no need for a separate reference list.
            """)

        else:
            system_message_prompt = SystemMessagePromptTemplate.from_template(
                """You are noteworthy standing scholar with extensive expertise in {expertise_areas}, you've demonstrated a unique knack for critically examining academic content. You will contribute to a narrative review paper titled, {subject}, using these talents. 
                
                For the final product, please be sure to cite the relevant papers without adding any formal reference list. We would also like to emphasize that the final output should be consolidated into one continuous paragraph.""")

            human_message_prompt = HumanMessagePromptTemplate.from_template(
                """weave the provided narratives - {narrative_content}, pertaining to a section titled {section}, into a singular, coherent, and structured passage. If the narratives are empty or do not align with the {section}, please return 'None'.
                
                It's crucial to ensure a smooth transition between ideas while maintaining consistency throughout the narrative. Furthermore, please make sure to accurately cite each source mentioned in these narratives following the APA style for in-text citation in the final combined narrative. We insist on the exclusion of a formal reference list in the final output, asking you to merely cite the papers.
                
                Also, it is imperative that the final output be a single comprehensive paragraph.
                """)

        chat_prompt = ChatPromptTemplate.from_messages(
            [system_message_prompt, human_message_prompt])

        return chat_prompt

    def section_writing(self, **kwargs):
        """write a section of a narrative review paper 

        Returns:
            str: section content
        """

        if len(kwargs.get('texts')) == 1:
            chat_prompt = self.section_writing_prompt(p_type='short')
            chain = LLMChain(llm=kwargs.get('llm_model'), prompt=chat_prompt)

            output = chain.run(expertise_areas=kwargs.get('expertise_areas'), subject=kwargs.get(
                'subject'), section=kwargs.get('section'), papers_summary=kwargs.get('texts')[0])
        else:
            print('many')
            narrative_content = []
            for i in range(len(kwargs.get('texts'))):
                chat_prompt = self.section_writing_prompt(p_type='short')
                chain = LLMChain(llm=kwargs.get(
                    'llm_model'), prompt=chat_prompt)
                r = chain.run(expertise_areas=kwargs.get('expertise_areas'), subject=kwargs.get(
                    'subject'), section=kwargs.get('section'), papers_summary=kwargs.get('texts')[i])
                narrative_content.append(r)

                time.sleep(kwargs.get('to_sleep'))

            chat_prompt = self.section_writing_prompt(p_type='combine')
            chain = LLMChain(llm=kwargs.get('llm_model'), prompt=chat_prompt)
            output = chain.run(expertise_areas=kwargs.get('expertise_areas'), subject=kwargs.get(
                'subject'), narrative_content=narrative_content, section=kwargs.get('section'))

        return output

    def references(self, **kwargs):
        """list references used in manuscripts 

        Returns:
            str: list of references
        """
        system_message_prompt = SystemMessagePromptTemplate.from_template(
            """you're an accomplished researcher skilled in formatting references according to the APA style.""")
        human_message_prompt = HumanMessagePromptTemplate.from_template(
            """provided these references {references_text}. I'd appreciate it if you could ensure they are correctly formatted according to the APA style guidelines. Just return the list of references without anything else""")
        chat_prompt = ChatPromptTemplate.from_messages(
            [system_message_prompt, human_message_prompt])

        chain = LLMChain(llm=kwargs.get(
            'llm_model'), prompt=chat_prompt)
        output = chain.run(references_text=kwargs.get('references_text'))

        return output
